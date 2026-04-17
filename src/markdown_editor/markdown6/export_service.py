"""Export service — thin adapter over `html_renderer_core`.

Supports PDF, DOCX, and HTML export. HTML always goes through the shared
`html_renderer_core.render_html_document` (decisions A, B1, H1 in
local/html-export-unify.md). This module contains zero custom markdown
logic, zero custom CSS, zero custom template — just export orchestration
(ephemeral ctx copy, title post-processing, file writing, pandoc/weasyprint
dispatch).
"""

import html as _html
import re
import subprocess
from pathlib import Path
from typing import TYPE_CHECKING

from markdown_editor.markdown6.logger import getLogger

if TYPE_CHECKING:
    from markdown_editor.markdown6.app_context import AppContext

logger = getLogger(__name__)


def has_pandoc() -> bool:
    """Check if pandoc is available on the system."""
    from markdown_editor.markdown6.tool_paths import has_pandoc as _has_pandoc
    return _has_pandoc()


# Markdown image syntax matching `![alt](path.dot)` — captures the
# path so we can rewrite it to absolute form when concatenating files
# from different source directories. URL-schemed and already-absolute
# paths are skipped in `_absolutize_dot_image_paths` below.
# IGNORECASE mirrors `GraphvizImagePostprocessor.DOT_IMAGE_PATTERN`
# so `.DOT` / `.Dot` / `.dot` are all rewritten consistently.
_DOT_IMAGE_RE = re.compile(
    r'(!\[[^\]]*\]\()([^)\s]+\.dot)(\s*(?:"[^"]*"|\'[^\']*\')?\s*\))',
    re.IGNORECASE,
)


def _absolutize_dot_image_paths(content: str, source_path: Path) -> str:
    """Rewrite relative `.dot` markdown image refs to absolute paths.

    When the combined export concatenates files from different parent
    directories, a single `graphviz_base_path` on the renderer can't
    resolve all relative `.dot` refs correctly — each file's `./x.dot`
    is anchored to its own parent. Rewriting to absolute paths
    pre-concatenation makes each ref self-resolving; the renderer's
    `GraphvizImagePostprocessor` handles absolute paths directly
    (absolute `Path(base) / absolute_path` returns the absolute path).

    Skips URLs and already-absolute paths. Scope: markdown image
    syntax `![alt](path.dot)` only — raw HTML `<img src="./x.dot">`
    in markdown is not pre-processed here (rare enough to leave
    to the renderer's postprocessor + graphviz_base_path).
    """
    base = source_path.parent.resolve()

    def _sub(m: re.Match) -> str:
        prefix, path, suffix = m.groups()
        if path.startswith(('http://', 'https://', '/')):
            return m.group(0)
        absolute = (base / path).resolve()
        return f'{prefix}{absolute}{suffix}'

    return _DOT_IMAGE_RE.sub(_sub, content)


def combine_project_markdown(
    documents: list[tuple[Path, str]],
    *,
    include_toc: bool = False,
    page_breaks: bool = False,
    output_format: str = "html",
) -> str:
    """Stitch multiple markdown documents into one combined source.

    The single source of truth for building a multi-file project
    export's combined markdown — shared by the CLI (`mde export -p`,
    `mde export a.md b.md c.md`) and the GUI project export dialog.

    Args:
        documents: List of `(source_path, markdown_content)` tuples.
            The path's stem is used to derive the display name and
            TOC anchor. The path's parent is also used to resolve
            each document's relative `.dot` image references to
            absolute paths before concatenation — without this the
            renderer's single `graphviz_base_path` couldn't
            correctly resolve all references when files come from
            different directories. Callers handle file I/O themselves
            (so they can drive progress dialogs or cancel logic
            independently).
        include_toc: If True, prepends a "Table of Contents" section
            with a numbered list linking to each document's heading
            anchor.
        page_breaks: If True, inserts a separator between documents
            (a `<div style="page-break-before: always;">` for HTML,
            a `\\n---\\n` horizontal rule for markdown).
        output_format: `"html"` or `"markdown"`. Controls the
            page-break markup style; the actual HTML rendering step
            happens downstream in `markdown_to_html`.

    Returns:
        The combined markdown source, ready to be handed to
        `markdown_to_html` (for HTML export) or written directly
        (for markdown export).

    Behavior: no page break is inserted before the first document
    (both with and without a TOC). This matches the CLI's historical
    behavior; the GUI dialog's pre-refactor code inserted a break
    before every document including the first, which was an
    inadvertent divergence fixed here.
    """
    parts: list[str] = []

    if include_toc:
        parts.append("# Table of Contents\n")
        for i, (path, _content) in enumerate(documents, 1):
            name = path.stem.replace("-", " ").replace("_", " ").title()
            anchor = name.lower().replace(" ", "-")
            parts.append(f"{i}. [{name}](#{anchor})")
        parts.append("\n---\n")

    for i, (path, content) in enumerate(documents):
        if page_breaks and i > 0:
            if output_format == "html":
                parts.append('<div style="page-break-before: always;"></div>\n')
            else:
                parts.append("\n---\n")
        parts.append(_absolutize_dot_image_paths(content, path))
        parts.append("\n\n")

    return "\n".join(parts)


def markdown_to_html(
    content: str,
    title: str = "Document",
    ctx: "AppContext | None" = None,
    source_path: "Path | None" = None,
) -> str:
    """Convert markdown content to a complete HTML document.

    Thin adapter over `html_renderer_core.render_html_document`. Every
    HTML export path — CLI, GUI single-file, GUI project, and
    WeasyPrint PDF — flows through this function.

    Args:
        content: Markdown source text.
        title: Document title — populates the `<title>` tag.
        ctx: AppContext to read settings from. Defaults to the global
            app context (`get_app_context()`). Never mutated — an
            ephemeral copy is made internally (decision E) and export-
            specific overrides are applied to the copy.
        source_path: If the caller knows the source markdown file's
            location, pass it so the renderer can resolve relative
            `.dot` image references (e.g. `![x](./graph.dot)`) against
            its parent directory. Optional; callers for whom there is
            no single source path (stdin, project exports that
            concatenate multiple files) pass None.

    Implementation notes:
        - `editor.scroll_past_end=False` is forced on the ephemeral
          copy so exports never emit the trailing 80vh placeholder
          (decision E).
        - `<title>` is populated via a `.replace()` post-processing
          step on the renderer's `<title></title>` placeholder. This
          is explicitly tech debt (decision T1) — the clean fix is to
          add `session.current_file_path` to AppContext so the
          renderer can read the title from ctx directly. See FUTURE
          WORK in local/html-export-unify.md.
    """
    # Imported here to keep the module import graph headless-friendly:
    # app_context and html_renderer_core pull in PySide6.QtCore at import
    # time, which we want to defer until actually needed.
    from markdown_editor.markdown6 import html_renderer_core
    from markdown_editor.markdown6.app_context import get_app_context

    if ctx is None:
        ctx = get_app_context()

    # Decision E — export path obtains an ephemeral copy of the caller's
    # ctx and mutates only the copy. The live/caller ctx is never changed.
    export_ctx = ctx.ephemeral_copy()
    export_ctx.set("editor.scroll_past_end", False)
    if source_path is not None:
        # Resolve to an absolute path so diagram resolution is
        # independent of the process's CWD (which may differ from the
        # source file's directory when `mde` is invoked via absolute
        # path from a different shell dir).
        export_ctx.set(
            "_render.graphviz_base_path",
            str(Path(source_path).resolve().parent),
        )

    html_doc = html_renderer_core.render_html_document(content, export_ctx)

    # TECH DEBT (decision T1): the renderer emits `<title></title>` empty
    # because AppContext doesn't know the current document path. This
    # .replace() is the interim wiring to populate the tag. Delete this
    # block and have the renderer read from `ctx.get("session.current_file_path")`
    # (or similar) once that setting exists — see FUTURE WORK in
    # local/html-export-unify.md.
    html_doc = html_doc.replace(
        "<title></title>",
        f"<title>{_html.escape(title)}</title>",
        1,
    )
    return html_doc


def export_html(
    content: str,
    output_path: str | Path,
    title: str = "Document",
    ctx: "AppContext | None" = None,
    source_path: "Path | None" = None,
) -> None:
    """Export markdown to an HTML file on disk."""
    html_doc = markdown_to_html(content, title=title, ctx=ctx, source_path=source_path)
    Path(output_path).write_text(html_doc, encoding="utf-8")


def export_pdf(
    content: str,
    output_path: str | Path,
    title: str = "Document",
    use_pandoc: bool = False,
) -> None:
    """Export markdown to PDF.

    Args:
        content: Markdown content to export
        output_path: Path to save the PDF
        title: Document title
        use_pandoc: If True and pandoc is available, use pandoc+xelatex
    """
    if use_pandoc and has_pandoc():
        _export_pdf_pandoc(content, output_path)
    else:
        _export_pdf_weasyprint(content, output_path, title)


def export_docx(
    content: str,
    output_path: str | Path,
    title: str = "Document",
    use_pandoc: bool = False,
) -> None:
    """Export markdown to DOCX.

    Args:
        content: Markdown content to export
        output_path: Path to save the DOCX
        title: Document title
        use_pandoc: If True and pandoc is available, use pandoc
    """
    if use_pandoc and has_pandoc():
        _export_docx_pandoc(content, output_path)
    else:
        _export_docx_python(content, output_path, title)


def _export_pdf_pandoc(content: str, output_path: str | Path) -> None:
    """Export to PDF using pandoc."""
    from markdown_editor.markdown6.temp_files import create_temp_file
    from markdown_editor.markdown6.tool_paths import get_pandoc_path

    temp_path = create_temp_file(suffix=".md", content=content)

    pandoc = get_pandoc_path() or "pandoc"
    cmd = [pandoc, str(temp_path), "-o", str(output_path), "--pdf-engine=xelatex"]
    logger.info(f"Running pandoc PDF export: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    if result.returncode != 0:
        logger.error(f"Pandoc PDF export failed (rc={result.returncode}): {result.stderr}")
        raise ExportError(f"Pandoc error: {result.stderr or 'Unknown error'}")


def _export_pdf_weasyprint(content: str, output_path: str | Path, title: str) -> None:
    """Export to PDF using weasyprint."""
    try:
        from weasyprint import HTML
    except ImportError:
        raise ExportError(
            "PDF export requires either pandoc or weasyprint.\n"
            "Install weasyprint with: pip install weasyprint"
        )

    html = markdown_to_html(content, title)
    logger.info(f"Exporting PDF via weasyprint to {output_path}")
    HTML(string=html).write_pdf(str(output_path))


def _export_docx_pandoc(content: str, output_path: str | Path) -> None:
    """Export to DOCX using pandoc."""
    from markdown_editor.markdown6.temp_files import create_temp_file
    from markdown_editor.markdown6.tool_paths import get_pandoc_path

    temp_path = create_temp_file(suffix=".md", content=content)

    pandoc = get_pandoc_path() or "pandoc"
    cmd = [pandoc, str(temp_path), "-o", str(output_path)]
    logger.info(f"Running pandoc DOCX export: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    if result.returncode != 0:
        logger.error(f"Pandoc DOCX export failed (rc={result.returncode}): {result.stderr}")
        raise ExportError(f"Pandoc error: {result.stderr or 'Unknown error'}")


def _export_docx_python(content: str, output_path: str | Path, title: str) -> None:
    """Export to DOCX using python-docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
        from docx.shared import Inches, Pt
    except ImportError:
        raise ExportError(
            "DOCX export requires either pandoc or python-docx.\n"
            "Install python-docx with: pip install python-docx"
        )

    doc = Document()

    # Parse markdown and convert to docx
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.5)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Headings
        if line.startswith('#'):
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break
            text = line[level:].strip()
            if level <= 9:
                doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraph - collect consecutive non-empty lines
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('```'):
            para_lines.append(lines[i])
            i += 1

        para_text = ' '.join(para_lines)

        # Basic inline formatting
        p = doc.add_paragraph()
        _add_formatted_text(p, para_text)

    doc.save(str(output_path))


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to paragraph with basic markdown formatting."""
    import re

    # Simple pattern matching for bold, italic, code
    # This is simplified - full markdown parsing would be more complex
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            paragraph.add_run(part)


class ExportError(Exception):
    """Raised when an export operation fails."""
    pass
